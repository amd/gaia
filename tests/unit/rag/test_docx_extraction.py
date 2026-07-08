# Copyright(C) 2025-2026 Advanced Micro Devices, Inc. All rights reserved.
# SPDX-License-Identifier: MIT
"""
Unit tests for Word (.docx) extraction in gaia.rag.sdk.

Covers:
- Paragraph text extraction
- Table cell text extraction (tables must not be dropped)
- Document-order interleaving of paragraphs and tables
- Routing through the _extract_text_from_file dispatcher
- Actionable errors for corrupted / non-.docx files

Fixtures are built programmatically with python-docx so the tests remain
hermetic and don't require committing binary fixture files to the repo.
"""

from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

docx = pytest.importorskip("docx")

from docx import Document  # noqa: E402
from docx.oxml import parse_xml  # noqa: E402
from docx.oxml.ns import nsdecls  # noqa: E402

from gaia.rag.sdk import RAGSDK, RAGConfig  # noqa: E402

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def rag(tmp_path: Path) -> RAGSDK:
    """A RAGSDK instance scoped to tmp_path with heavy ML deps stubbed out.

    Same pattern as test_pptx_extraction.py — stubs sentence-transformers,
    faiss, VLM, and chat/LLM initialization.
    """
    config = RAGConfig(
        cache_dir=str(tmp_path / ".gaia"),
        show_stats=False,
        use_local_llm=False,
    )

    fake_vlm = MagicMock(name="VLMClient")
    fake_vlm.check_availability.return_value = False

    with (
        patch.object(RAGSDK, "_check_dependencies", return_value=None),
        patch("gaia.rag.sdk.AgentSDK", autospec=True) as mock_agent_sdk,
        patch("gaia.llm.VLMClient", return_value=fake_vlm),
    ):
        mock_agent_sdk.return_value = MagicMock(name="AgentSDK")
        instance = RAGSDK(config=config)

    yield instance


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _create_docx(path: Path, paragraphs: list, table: list | None = None) -> None:
    """Create a test .docx programmatically.

    *paragraphs* is a list of paragraph strings. *table*, if given, is a
    list of rows (each a list of cell strings); the table is appended after
    the paragraphs.
    """
    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)

    if table:
        num_rows = len(table)
        num_cols = len(table[0]) if table else 0
        tbl = doc.add_table(rows=num_rows, cols=num_cols)
        for r_idx, row in enumerate(table):
            for c_idx, cell_text in enumerate(row):
                tbl.cell(r_idx, c_idx).text = cell_text

    doc.save(str(path))


# ---------------------------------------------------------------------------
# Tests — Text Extraction
# ---------------------------------------------------------------------------


class TestDocxTextExtraction:
    """Tests for _extract_text_from_docx text handling."""

    def test_basic_paragraph_extraction(self, rag, tmp_path):
        """Paragraph text is extracted."""
        docx_path = tmp_path / "basic.docx"
        _create_docx(docx_path, ["Hello World", "This is the second paragraph."])

        text = rag._extract_text_from_docx(str(docx_path))

        assert "Hello World" in text
        assert "This is the second paragraph." in text

    def test_table_cells_extracted(self, rag, tmp_path):
        """Table cell text is extracted — tables must not be dropped."""
        docx_path = tmp_path / "table.docx"
        _create_docx(
            docx_path,
            ["Intro paragraph"],
            table=[
                ["Name", "Age", "City"],
                ["Alice", "30", "Boston"],
                ["Bob", "25", "Denver"],
            ],
        )

        text = rag._extract_text_from_docx(str(docx_path))

        # Paragraph and every table cell appear
        assert "Intro paragraph" in text
        for cell in ("Name", "Age", "City", "Alice", "30", "Boston", "Bob", "Denver"):
            assert cell in text

    def test_document_order_preserved(self, rag, tmp_path):
        """Paragraphs and tables stay in document order (table interleaved)."""
        docx_path = tmp_path / "order.docx"
        _create_docx(
            docx_path,
            ["BEFORE_TABLE"],
            table=[["CELL_VALUE"]],
        )
        # Append a paragraph after the table by re-opening
        doc = Document(str(docx_path))
        doc.add_paragraph("AFTER_TABLE")
        doc.save(str(docx_path))

        text = rag._extract_text_from_docx(str(docx_path))

        assert text.index("BEFORE_TABLE") < text.index("CELL_VALUE")
        assert text.index("CELL_VALUE") < text.index("AFTER_TABLE")

    def test_empty_paragraphs_skipped(self, rag, tmp_path):
        """Blank paragraphs do not add empty lines/noise."""
        docx_path = tmp_path / "blanks.docx"
        _create_docx(docx_path, ["Real content", "", "   ", "More content"])

        text = rag._extract_text_from_docx(str(docx_path))

        assert "Real content" in text
        assert "More content" in text
        # No run of blank lines
        assert "\n\n\n" not in text


# ---------------------------------------------------------------------------
# Tests — Rich content (content controls, nested tables, hyperlinks)
# ---------------------------------------------------------------------------


class TestDocxRichContent:
    """Form/template documents put real data in content controls and nested
    tables — extracting them is the difference between indexing the labels
    and indexing the answers."""

    def test_inline_content_control_captured(self, rag, tmp_path):
        """Text in an inline content control (w:sdt) is captured, not dropped.

        Paragraph.text only sees direct runs; a filled form field lives in a
        nested w:sdt run that would otherwise be silently lost.
        """
        doc = Document()
        para = doc.add_paragraph("Name: ")
        sdt = parse_xml(
            f"<w:sdt {nsdecls('w')}><w:sdtContent><w:r><w:t>John Smith</w:t>"
            f"</w:r></w:sdtContent></w:sdt>"
        )
        para._p.append(sdt)
        docx_path = tmp_path / "inline_cc.docx"
        doc.save(str(docx_path))

        text = rag._extract_text_from_docx(str(docx_path))

        assert "Name:" in text
        assert "John Smith" in text

    def test_block_content_control_captured(self, rag, tmp_path):
        """A block-level content control's paragraphs are recursed into."""
        doc = Document()
        doc.add_paragraph("Before")
        sdt = parse_xml(
            f"<w:sdt {nsdecls('w')}><w:sdtContent>"
            f"<w:p><w:r><w:t>Inside content control</w:t></w:r></w:p>"
            f"</w:sdtContent></w:sdt>"
        )
        doc.element.body.append(sdt)
        docx_path = tmp_path / "block_cc.docx"
        doc.save(str(docx_path))

        text = rag._extract_text_from_docx(str(docx_path))

        assert "Before" in text
        assert "Inside content control" in text

    def test_nested_table_in_cell_captured(self, rag, tmp_path):
        """A table nested inside a table cell is not dropped."""
        doc = Document()
        outer = doc.add_table(rows=1, cols=1)
        cell = outer.cell(0, 0)
        cell.text = "Outer cell"
        nested = cell.add_table(rows=1, cols=2)
        nested.cell(0, 0).text = "NestedA"
        nested.cell(0, 1).text = "NestedB"
        docx_path = tmp_path / "nested_table.docx"
        doc.save(str(docx_path))

        text = rag._extract_text_from_docx(str(docx_path))

        assert "Outer cell" in text
        assert "NestedA" in text
        assert "NestedB" in text

    def test_hyperlink_text_captured(self, rag, tmp_path):
        """Hyperlink anchor text (a w:hyperlink-wrapped run) is captured."""
        doc = Document()
        para = doc.add_paragraph("See ")
        link = parse_xml(
            f'<w:hyperlink {nsdecls("w", "r")} r:id="rId99">'
            f"<w:r><w:t>the documentation</w:t></w:r></w:hyperlink>"
        )
        para._p.append(link)
        docx_path = tmp_path / "link.docx"
        doc.save(str(docx_path))

        text = rag._extract_text_from_docx(str(docx_path))

        assert "the documentation" in text

    def test_tabs_and_breaks_become_whitespace(self, rag, tmp_path):
        """w:tab / w:br must inject whitespace so adjacent runs don't glue.

        A naive ``w:t``-only join turns a tab-separated label/value pair into
        an unsearchable concatenation (``Column1Column2``).
        """
        doc = Document()
        para = doc.add_paragraph()
        para._p.append(
            parse_xml(
                f"<w:r {nsdecls('w')}><w:t>Column1</w:t><w:tab/>"
                f"<w:t>Column2</w:t><w:br/><w:t>NextLine</w:t></w:r>"
            )
        )
        docx_path = tmp_path / "tabs.docx"
        doc.save(str(docx_path))

        text = rag._extract_text_from_docx(str(docx_path))

        assert "Column1Column2" not in text
        assert "Column1" in text and "Column2" in text and "NextLine" in text

    def test_intra_word_runs_not_split(self, rag, tmp_path):
        """Runs split mid-word (formatting boundaries) must NOT gain spaces."""
        doc = Document()
        para = doc.add_paragraph()
        # "Hel" + "lo Wor" + "ld" — formatting splits, not word boundaries.
        para._p.append(parse_xml(f"<w:r {nsdecls('w')}><w:t>Hel</w:t></w:r>"))
        para._p.append(
            parse_xml(
                f'<w:r {nsdecls("w")}><w:t xml:space="preserve">lo Wor</w:t></w:r>'
            )
        )
        para._p.append(parse_xml(f"<w:r {nsdecls('w')}><w:t>ld</w:t></w:r>"))
        docx_path = tmp_path / "intraword.docx"
        doc.save(str(docx_path))

        text = rag._extract_text_from_docx(str(docx_path))

        assert "Hello World" in text

    def test_textbox_text_captured_once(self, rag, tmp_path):
        """DrawingML textbox text is captured exactly once (mc:Fallback skipped).

        Word writes a shape as mc:AlternateContent with a DrawingML mc:Choice
        and a VML mc:Fallback that carries the SAME text — a recursive w:t walk
        would otherwise emit it twice and produce garbage duplicate tokens.
        """
        doc = Document()
        para = doc.add_paragraph("Body before. ")
        # Minimal AlternateContent: Choice (DrawingML) + Fallback (VML) both
        # carry the same text, exactly as Word writes a textbox/shape. The
        # walker must descend into Choice and skip Fallback.
        ns = (
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"'
        )
        alt = parse_xml(
            f"<w:r {ns}><mc:AlternateContent>"
            f'<mc:Choice Requires="wps">'
            f"<w:p><w:r><w:t>TEXTBOX_CONTENT</w:t></w:r></w:p></mc:Choice>"
            f"<mc:Fallback>"
            f"<w:p><w:r><w:t>TEXTBOX_CONTENT</w:t></w:r></w:p></mc:Fallback>"
            f"</mc:AlternateContent></w:r>"
        )
        para._p.append(alt)
        docx_path = tmp_path / "textbox.docx"
        doc.save(str(docx_path))

        text = rag._extract_text_from_docx(str(docx_path))

        assert text.count("TEXTBOX_CONTENT") == 1
        assert "Body before." in text

    def test_sdt_wrapped_table_row_captured(self, rag, tmp_path):
        """A row wrapped in a repeating-section content control is not dropped."""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "PlainRow"
        # Wrap a second row in a w:sdt directly under the table element.
        sdt_row = parse_xml(
            f"<w:sdt {nsdecls('w')}><w:sdtContent>"
            f"<w:tr><w:tc><w:p><w:r><w:t>WRAPPED_ROW</w:t></w:r></w:p></w:tc></w:tr>"
            f"</w:sdtContent></w:sdt>"
        )
        table._tbl.append(sdt_row)
        docx_path = tmp_path / "sdt_row.docx"
        doc.save(str(docx_path))

        text = rag._extract_text_from_docx(str(docx_path))

        assert "PlainRow" in text
        assert "WRAPPED_ROW" in text


# ---------------------------------------------------------------------------
# Tests — Edge Cases / Errors
# ---------------------------------------------------------------------------


class TestDocxErrors:
    """Tests for corrupted / invalid .docx handling."""

    def test_corrupted_file_raises(self, rag, tmp_path):
        """Garbage bytes with a .docx extension raise an actionable ValueError."""
        docx_path = tmp_path / "corrupt.docx"
        docx_path.write_bytes(b"this is not a docx file at all")

        with pytest.raises(ValueError, match="corrupted|not a valid"):
            rag._extract_text_from_docx(str(docx_path))

    def test_error_names_the_file(self, rag, tmp_path):
        """The actionable error names the offending file."""
        docx_path = tmp_path / "broken_report.docx"
        docx_path.write_bytes(b"PK\x03\x04 not really a zip body")

        with pytest.raises(ValueError, match="broken_report.docx"):
            rag._extract_text_from_docx(str(docx_path))


# ---------------------------------------------------------------------------
# Tests — Dispatcher integration
# ---------------------------------------------------------------------------


class TestDocxDispatcher:
    """Tests for .docx routing through _extract_text_from_file."""

    def test_docx_dispatches_correctly(self, rag, tmp_path):
        """_extract_text_from_file routes .docx to _extract_text_from_docx."""
        docx_path = tmp_path / "dispatch.docx"
        _create_docx(
            docx_path,
            ["Dispatch Test"],
            table=[["Works", "Fine"]],
        )

        text, metadata = rag._extract_text_from_file(str(docx_path))

        assert "Dispatch Test" in text
        assert "Works" in text
        assert "Fine" in text
        # Word has no page concept exposed by python-docx
        assert metadata["num_pages"] is None
